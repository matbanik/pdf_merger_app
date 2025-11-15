import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, simpledialog
import os
import json
import threading
import time
import fitz  # PyMuPDF
import re
import tempfile
import multiprocessing
import subprocess
import shutil
import webbrowser
# Marker imports moved to functions to allow environment variable setting first
import logging

# --- Constants and Global Variables ---
SETTINGS_FILE = "settings.json"
DEFAULT_OUTPUT_FILENAME = "MergedPDFs.pdf"
# Determine the user's Downloads directory across different OS
DOWNLOADS_PATH = os.path.join(os.path.expanduser("~"), "Downloads")
# Local models directory in app folder
MODELS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "models")

# Global variables for controlling the merge process thread
merge_thread = None
merge_stop_event = threading.Event()
merge_pause_event = threading.Event()
merge_running = False
merge_paused = False

class PDFMergerApp:
    def __init__(self, master):
        self.master = master
        master.title("Document Merger & PII Scrubber - Multi-Format Support")
        master.geometry("800x1000") # Increased height for all sections

        self.pdf_files = [] # List to store full paths of all supported files
        self.input_folder = DOWNLOADS_PATH # Default input folder
        self.output_folder = DOWNLOADS_PATH # Default output folder
        self.total_word_count = 0 # Accumulator for total words

        # --- Configuration Variables ---
        self.remove_timestamps_var = tk.BooleanVar(value=False)
        self.remove_images_var = tk.BooleanVar(value=False)
        self.remove_pii_var = tk.BooleanVar(value=False)
        self.custom_pii_var = tk.StringVar(value="")
        # New: Variables for splitting output
        self.split_by_words_var = tk.BooleanVar(value=False)
        self.split_word_count_var = tk.StringVar(value="10000")
        # New: Variable for markdown output
        self.generate_markdown_var = tk.BooleanVar(value=False)
        # New: Variable for simple markdown (without OCR)
        self.simple_markdown_var = tk.BooleanVar(value=False)
        # New: Variable for markdown type (radio button)
        self.markdown_type_var = tk.StringVar(value="simple")  # "simple" or "advanced"
        # New: Variable for GPU acceleration
        self.use_gpu_var = tk.BooleanVar(value=False)
        # New: Variable for models directory
        self.models_directory = MODELS_DIR  # Default to app directory
        # New: Variable for qpdf executable path
        self.qpdf_path = None  # Will be loaded from settings
        # New: Console visibility and filter variables
        self.console_visible_var = tk.BooleanVar(value=True)
        self.console_filter_level_var = tk.StringVar(value="ALL")
        # Message buffer for filtering
        self.console_message_buffer = []  # List of (message, tag, level) tuples
        # New: Multi-format support variables
        self.output_file_type_var = tk.StringVar(value="PDF")  # Output format
        self.output_filename_var = tk.StringVar(value="")  # Optional custom output filename
        self.preserve_formatting_var = tk.BooleanVar(value=False)  # Preserve formatting when possible

        # Initialize widgets first so console_output exists before load_settings
        self.create_widgets() # Build the GUI elements
        self.load_settings() # Load saved settings on startup
        self.update_word_count_display() # Update the word count label initially
        self._update_pii_field_visibility() # Set initial state of custom PII field
        self._update_split_field_visibility() # Set initial state of split field

    def create_widgets(self):
        """Creates and lays out all the GUI widgets."""
        # --- Top Section: Word Count ---
        top_frame = tk.Frame(self.master, bd=2, relief="groove", padx=10, pady=10)
        top_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=10)

        self.total_words_label = tk.Label(top_frame, text=f"Total Words: {self.total_word_count}", font=("Arial", 14, "bold"))
        self.total_words_label.pack(side=tk.TOP, pady=5)

        # --- INPUT Configuration Section ---
        input_config_frame = tk.LabelFrame(self.master, text="INPUT Configuration", bd=2, relief="groove", padx=10, pady=10)
        input_config_frame.pack(side=tk.TOP, fill=tk.BOTH, expand=True, padx=10, pady=5)

        # Input Folder row
        input_folder_frame = tk.Frame(input_config_frame)
        input_folder_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 10))

        tk.Label(input_folder_frame, text="Input Folder:").pack(side=tk.LEFT)
        self.input_folder_label = tk.Label(input_folder_frame, text=self.input_folder, bg="lightgray", anchor="w", relief="sunken")
        self.input_folder_label.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        self.select_input_folder_btn = tk.Button(input_folder_frame, text="Select", command=self.select_input_folder)
        self.select_input_folder_btn.pack(side=tk.LEFT, padx=2)

        # Separator
        separator_input = tk.Frame(input_folder_frame, width=2, bg="gray", relief=tk.SUNKEN)
        separator_input.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # Add File(s) button
        self.add_btn = tk.Button(input_folder_frame, text="Add File(s)", command=self.add_pdf_file, width=12)
        self.add_btn.pack(side=tk.LEFT, padx=2)

        # Files for Merger list
        tk.Label(input_config_frame, text="Files for Merger:", font=("Arial", 12)).pack(side=tk.TOP, anchor="w", pady=(5,2))

        list_container = tk.Frame(input_config_frame)
        list_container.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        self.pdf_listbox = tk.Listbox(list_container, selectmode=tk.EXTENDED, height=8)
        self.pdf_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 10))
        self.pdf_listbox.bind("<<ListboxSelect>>", self.on_listbox_select)

        scrollbar = tk.Scrollbar(list_container, orient="vertical", command=self.pdf_listbox.yview)
        scrollbar.pack(side=tk.LEFT, fill=tk.Y)
        self.pdf_listbox.config(yscrollcommand=scrollbar.set)

        # Control Buttons Section (toolbar below the list)
        control_frame = tk.Frame(input_config_frame)
        control_frame.pack(side=tk.TOP, fill=tk.X, pady=(10, 0))

        # Toolbar with icon buttons
        toolbar = tk.Frame(control_frame)
        toolbar.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Play button (Start Merge)
        self.start_btn = tk.Button(
            toolbar,
            text="▶",
            font=("Arial", 14, "bold"),
            command=self.start_merge,
            width=3,
            relief=tk.RAISED,
            bd=1
        )
        self.start_btn.pack(side=tk.LEFT, padx=2)
        self.start_btn.bind("<Enter>", lambda e: self.start_btn.config(cursor="hand2"))
        self.start_btn.bind("<Leave>", lambda e: self.start_btn.config(cursor=""))

        # Stop button (Stop Merge)
        self.stop_btn = tk.Button(
            toolbar,
            text="■",
            font=("Arial", 14, "bold"),
            command=self.stop_merge,
            width=3,
            state=tk.DISABLED,
            relief=tk.RAISED,
            bd=1
        )
        self.stop_btn.pack(side=tk.LEFT, padx=2)
        self.stop_btn.bind("<Enter>", lambda e: self.stop_btn.config(cursor="hand2") if self.stop_btn.cget("state") == tk.NORMAL else None)

        # Pause button (using double bar symbol)
        self.pause_btn = tk.Button(
            toolbar,
            text="⏸",
            font=("Arial", 14, "bold"),
            command=self.pause_merge,
            width=3,
            state=tk.DISABLED,
            relief=tk.RAISED,
            bd=1
        )
        self.pause_btn.pack(side=tk.LEFT, padx=2)
        self.pause_btn.bind("<Enter>", lambda e: self.pause_btn.config(cursor="hand2") if self.pause_btn.cget("state") == tk.NORMAL else None)

        # Separator 1
        separator1 = tk.Frame(toolbar, width=2, bg="gray", relief=tk.SUNKEN)
        separator1.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # Move Up button (↑)
        self.move_up_btn = tk.Button(
            toolbar,
            text="↑",
            font=("Arial", 14, "bold"),
            command=self.move_pdf_up,
            state=tk.DISABLED,
            width=3,
            relief=tk.RAISED,
            bd=1
        )
        self.move_up_btn.pack(side=tk.LEFT, padx=2)

        # Move Down button (↓)
        self.move_down_btn = tk.Button(
            toolbar,
            text="↓",
            font=("Arial", 14, "bold"),
            command=self.move_pdf_down,
            state=tk.DISABLED,
            width=3,
            relief=tk.RAISED,
            bd=1
        )
        self.move_down_btn.pack(side=tk.LEFT, padx=2)

        # Move to Top button (⇈)
        self.move_to_top_btn = tk.Button(
            toolbar,
            text="⇈",
            font=("Arial", 14, "bold"),
            command=self.move_pdf_to_top,
            state=tk.DISABLED,
            width=3,
            relief=tk.RAISED,
            bd=1
        )
        self.move_to_top_btn.pack(side=tk.LEFT, padx=2)

        # Move to Bottom button (⇊)
        self.move_to_bottom_btn = tk.Button(
            toolbar,
            text="⇊",
            font=("Arial", 14, "bold"),
            command=self.move_pdf_to_bottom,
            state=tk.DISABLED,
            width=3,
            relief=tk.RAISED,
            bd=1
        )
        self.move_to_bottom_btn.pack(side=tk.LEFT, padx=2)

        # Separator 2
        separator2 = tk.Frame(toolbar, width=2, bg="gray", relief=tk.SUNKEN)
        separator2.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=2)

        # Remove Selected button (text button)
        self.remove_btn = tk.Button(toolbar, text="Remove Selected", command=self.remove_pdf_file, state=tk.DISABLED, width=15)
        self.remove_btn.pack(side=tk.LEFT, padx=2)

        # Clear All button (text button)
        self.clear_all_btn = tk.Button(toolbar, text="Clear All", command=self.clear_all_pdfs, state=tk.DISABLED, width=12)
        self.clear_all_btn.pack(side=tk.LEFT, padx=2)

        # --- Tools Section ---
        tools_frame = tk.LabelFrame(self.master, text="Tools", bd=2, relief="groove", padx=10, pady=10)
        tools_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=5)

        # PDF Decryption
        decrypt_frame = tk.Frame(tools_frame)
        decrypt_frame.pack(side=tk.TOP, fill=tk.X)

        tk.Label(decrypt_frame, text="PDF Decryption (qpdf):", font=("Arial", 12)).pack(side=tk.LEFT, padx=5)
        # Dynamic button: "Locate qpdf" when not configured, "Decrypt PDF" when configured
        self.decrypt_btn = tk.Button(decrypt_frame, text="Locate qpdf", command=self._decrypt_or_locate, width=15)
        self.decrypt_btn.pack(side=tk.LEFT, padx=5, pady=5)

        # Display current qpdf path (if configured)
        self.qpdf_path_label = tk.Label(decrypt_frame, text="", fg="green", font=("Arial", 8))
        self.qpdf_path_label.pack(side=tk.LEFT, padx=5, pady=5)

        # Add link to download qpdf
        link_label = tk.Label(
            decrypt_frame,
            text="Get qpdf for Decryption",
            fg="blue",
            cursor="hand2",
            font=("Arial", 9, "underline")
        )
        link_label.pack(side=tk.RIGHT, padx=5)
        link_label.bind("<Button-1>", lambda e: self._open_qpdf_download_page())

        # --- OUTPUT Configuration Section ---
        output_config_frame = tk.LabelFrame(self.master, text="OUTPUT Configuration", bd=2, relief="groove", padx=10, pady=10)
        output_config_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=5)

        # Output File Type dropdown (at top)
        output_type_frame = tk.Frame(output_config_frame)
        output_type_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 5))

        tk.Label(output_type_frame, text="Select output file type:").pack(side=tk.LEFT)
        output_types = ["PDF", "ODT", "DOCX", "TXT", "RTF", "EPUB", "MD"]
        self.output_type_dropdown = tk.OptionMenu(output_type_frame, self.output_file_type_var, *output_types, command=self.on_output_type_change)
        self.output_type_dropdown.pack(side=tk.LEFT, padx=5)

        # Optional output filename
        output_filename_frame = tk.Frame(output_config_frame)
        output_filename_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 5))

        tk.Label(output_filename_frame, text="Output filename (optional):").pack(side=tk.LEFT)
        self.output_filename_entry = tk.Entry(output_filename_frame, textvariable=self.output_filename_var, width=30)
        self.output_filename_entry.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)
        self.output_filename_var.trace_add("write", lambda *args: self.save_settings())

        # Output Folder
        output_folder_frame = tk.Frame(output_config_frame)
        output_folder_frame.pack(side=tk.TOP, fill=tk.X, pady=(0, 10))

        tk.Label(output_folder_frame, text="Output Folder:").pack(side=tk.LEFT)
        self.output_folder_label = tk.Label(output_folder_frame, text=self.output_folder, bg="lightgray", anchor="w", relief="sunken")
        self.output_folder_label.pack(side=tk.LEFT, expand=True, fill=tk.X, padx=5)
        self.select_folder_btn = tk.Button(output_folder_frame, text="Select", command=self.select_output_folder)
        self.select_folder_btn.pack(side=tk.RIGHT)

        # Create two-column layout for options
        config_columns = tk.Frame(output_config_frame)
        config_columns.pack(side=tk.TOP, fill=tk.BOTH, expand=True)

        # --- Left Column: Basic Options (50%) ---
        left_column = tk.Frame(config_columns)
        left_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 5))

        # Preserve Formatting checkbox
        self.preserve_formatting_checkbox = tk.Checkbutton(left_column, text="Preserve Formatting (when possible)", variable=self.preserve_formatting_var, command=lambda: self.log_and_save_setting("Preserve Formatting", self.preserve_formatting_var), state=tk.DISABLED)
        self.preserve_formatting_checkbox.pack(anchor="w", padx=5, pady=2)

        self.remove_timestamps_checkbox = tk.Checkbutton(left_column, text="Remove Timestamps", variable=self.remove_timestamps_var, command=lambda: self.log_and_save_setting("Timestamps", self.remove_timestamps_var))
        self.remove_timestamps_checkbox.pack(anchor="w", padx=5, pady=2)

        self.remove_images_checkbox = tk.Checkbutton(left_column, text="Remove Images (extract text only)", variable=self.remove_images_var, command=lambda: self.log_and_save_setting("Images", self.remove_images_var))
        self.remove_images_checkbox.pack(anchor="w", padx=5, pady=2)

        self.remove_pii_checkbox = tk.Checkbutton(left_column, text="Remove PII (Names, Addresses, etc.)", variable=self.remove_pii_var, command=self.on_pii_checkbox_change)
        self.remove_pii_checkbox.pack(anchor="w", padx=5, pady=2)

        self.custom_pii_label = tk.Label(left_column, text="Custom Strings to Remove (comma-separated):")
        self.custom_pii_label.pack(anchor="w", padx=25, pady=(5,0))
        self.custom_pii_entry = tk.Entry(left_column, textvariable=self.custom_pii_var)
        self.custom_pii_entry.pack(fill=tk.X, padx=25, pady=2)
        self.custom_pii_var.trace_add("write", lambda *args: self.save_settings())

        # --- Right Column: Advanced Options (50%) ---
        right_column = tk.Frame(config_columns)
        right_column.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(5, 0))

        # Split by words widgets
        self.split_by_words_checkbox = tk.Checkbutton(right_column, text="Split output by words", variable=self.split_by_words_var, command=self.on_split_checkbox_change)
        self.split_by_words_checkbox.pack(anchor="w", padx=5, pady=(0,2))

        self.split_word_count_label = tk.Label(right_column, text="Number of words per file:")
        self.split_word_count_label.pack(anchor="w", padx=25, pady=(5,0))
        self.split_word_count_entry = tk.Entry(right_column, textvariable=self.split_word_count_var)
        self.split_word_count_entry.pack(fill=tk.X, padx=25, pady=2)
        self.split_word_count_var.trace_add("write", lambda *args: self.save_settings())

        # Markdown Options label and frame
        self.markdown_options_label = tk.Label(right_column, text="Markdown Options (.md):", font=("Arial", 10, "bold"), state=tk.DISABLED)
        self.markdown_options_label.pack(anchor="w", padx=5, pady=(10,5))

        markdown_options_frame = tk.Frame(right_column)
        markdown_options_frame.pack(fill=tk.X, padx=5, pady=2)

        # Radio buttons for markdown type
        self.simple_markdown_radio = tk.Radiobutton(
            markdown_options_frame,
            text="Simple Markdown (fast, no OCR)",
            variable=self.markdown_type_var,
            value="simple",
            command=self.on_markdown_type_change,
            state=tk.DISABLED
        )
        self.simple_markdown_radio.pack(anchor="w", padx=20, pady=2)

        self.advanced_markdown_radio = tk.Radiobutton(
            markdown_options_frame,
            text="Advanced Markdown (with OCR)",
            variable=self.markdown_type_var,
            value="advanced",
            command=self.on_markdown_type_change,
            state=tk.DISABLED
        )
        self.advanced_markdown_radio.pack(anchor="w", padx=20, pady=2)

        # GPU checkbox (under Advanced Markdown)
        self.use_gpu_checkbox = tk.Checkbutton(
            markdown_options_frame,
            text="Use GPU acceleration (if available)",
            variable=self.use_gpu_var,
            command=self.on_gpu_checkbox_change,
            state=tk.DISABLED
        )
        self.use_gpu_checkbox.pack(anchor="w", padx=40, pady=2)

        # Select Models Dir button (under Advanced Markdown)
        models_btn_frame = tk.Frame(markdown_options_frame)
        models_btn_frame.pack(fill=tk.X, padx=40, pady=2)

        self.preload_models_btn = tk.Button(models_btn_frame, text="Select Models Dir", command=self.preload_marker_models, width=15, state=tk.DISABLED)
        self.preload_models_btn.pack(side=tk.LEFT)

        # Models path label
        self.models_path_label = tk.Label(markdown_options_frame, text="", fg="green", font=("Arial", 9))
        self.models_path_label.pack(anchor="w", padx=40, pady=(2,5))

        # --- Console Output Section ---
        console_frame = tk.LabelFrame(self.master, text="Console Output", bd=2, relief="groove", padx=10, pady=10)
        console_frame.pack(side=tk.BOTTOM, fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Console header with controls
        console_header = tk.Frame(console_frame)
        console_header.pack(side=tk.TOP, fill=tk.X, pady=(0, 5))

        # Show/Hide checkbox
        self.console_show_checkbox = tk.Checkbutton(
            console_header,
            text="Show Console",
            variable=self.console_visible_var,
            command=self._toggle_console_visibility
        )
        self.console_show_checkbox.pack(side=tk.LEFT, padx=(0, 10))

        # Filter dropdown
        tk.Label(console_header, text="Filter:", font=("Arial", 10)).pack(side=tk.LEFT, padx=(0, 5))
        filter_levels = ["ALL", "DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"]
        self.console_filter_dropdown = tk.OptionMenu(console_header, self.console_filter_level_var, *filter_levels, command=self._on_filter_change)
        self.console_filter_dropdown.pack(side=tk.LEFT)

        self.console_output = scrolledtext.ScrolledText(console_frame, wrap=tk.WORD, height=8, bg="black", fg="lime", font=("Consolas", 10))
        self.console_output.pack(fill=tk.BOTH, expand=True)
        self.console_output.tag_config("info", foreground="white")
        self.console_output.tag_config("error", foreground="red")
        self.console_output.tag_config("progress", foreground="cyan")
        self.console_output.tag_config("success", foreground="green")
        self.console_output.tag_config("warning", foreground="yellow")
        self.console_output.tag_config("debug", foreground="gray")

        self.print_to_console("Welcome to Document Merger & PII Scrubber - Multi-Format Support!", "info")
        self.print_to_console("Select files (PDF, ODT, DOCX, TXT, RTF, EPUB, MD) and click 'Start Merge'.", "info")
        self.print_to_console(f"Default output folder: {self.output_folder}", "info")

        # Check qpdf availability on startup and update UI
        self._update_qpdf_ui_status()

        # Check GPU availability
        try:
            import torch
            if torch.cuda.is_available():
                gpu_name = torch.cuda.get_device_name(0)
                self.print_to_console(f"[INFO] GPU detected: {gpu_name}", "info")
            else:
                self.print_to_console("[INFO] No GPU detected, CPU processing available", "info")
        except ImportError:
            self.print_to_console("[INFO] PyTorch not available for GPU detection", "info")

    def select_input_folder(self):
        """Select the input folder for file browsing."""
        folder = filedialog.askdirectory(initialdir=self.input_folder, title="Select Input Folder")
        if folder:
            self.input_folder = folder
            self.input_folder_label.config(text=folder)
            self.print_to_console(f"Input folder set to: {folder}", "info")
            self.save_settings()

    def on_output_type_change(self, *args):
        """Handle output file type change."""
        output_type = self.output_file_type_var.get()
        self.print_to_console(f"Output type changed to: {output_type}", "info")

        # Enable/disable Markdown Options based on output type
        if output_type == "MD":
            self.markdown_options_label.config(state=tk.NORMAL)
            self.simple_markdown_radio.config(state=tk.NORMAL)
            self.advanced_markdown_radio.config(state=tk.NORMAL)
            self._update_markdown_controls_state()
        else:
            self.markdown_options_label.config(state=tk.DISABLED)
            self.simple_markdown_radio.config(state=tk.DISABLED)
            self.advanced_markdown_radio.config(state=tk.DISABLED)
            self.use_gpu_checkbox.config(state=tk.DISABLED)
            self.preload_models_btn.config(state=tk.DISABLED)

        self.save_settings()

    def on_markdown_type_change(self):
        """Handle markdown type radio button change."""
        self._update_markdown_controls_state()
        self.save_settings()

    def _update_markdown_controls_state(self):
        """Enable/disable markdown controls based on type selection."""
        if self.output_file_type_var.get() == "MD":
            if self.markdown_type_var.get() == "advanced":
                self.use_gpu_checkbox.config(state=tk.NORMAL)
                self.preload_models_btn.config(state=tk.NORMAL)
            else:
                self.use_gpu_checkbox.config(state=tk.DISABLED)
                self.preload_models_btn.config(state=tk.DISABLED)

    def on_pii_checkbox_change(self):
        """Handles changes to the PII checkbox state."""
        self.log_and_save_setting("PII", self.remove_pii_var)
        self._update_pii_field_visibility()

    def _update_pii_field_visibility(self):
        """Enables or disables the custom PII field based on the checkbox."""
        state = tk.NORMAL if self.remove_pii_var.get() else tk.DISABLED
        self.custom_pii_label.config(state=state)
        self.custom_pii_entry.config(state=state)

    # New: Method to handle split checkbox changes
    def on_split_checkbox_change(self):
        """Handles changes to the 'Split by words' checkbox state."""
        self.log_and_save_setting("Split by Words", self.split_by_words_var)
        self._update_split_field_visibility()

    # New: Method to enable/disable split word count field
    def _update_split_field_visibility(self):
        """Enables or disables the split word count field based on the checkbox."""
        state = tk.NORMAL if self.split_by_words_var.get() else tk.DISABLED
        self.split_word_count_label.config(state=state)
        self.split_word_count_entry.config(state=state)

    # New: Method to handle markdown checkbox changes
    def on_markdown_checkbox_change(self):
        """Handles changes to the 'Generate Markdown' checkbox state."""
        self.log_and_save_setting("Generate Markdown", self.generate_markdown_var)
        self._update_markdown_controls_visibility()
    
    # New: Method to handle simple markdown checkbox changes
    def on_simple_markdown_checkbox_change(self):
        """Handles changes to the 'Simple Markdown' checkbox state."""
        self.log_and_save_setting("Simple Markdown", self.simple_markdown_var)
        self._update_markdown_controls_visibility(log_message=True)
    
    def _update_markdown_controls_visibility(self, log_message=False):
        """Enables/disables GPU and models controls based on simple markdown setting."""
        if self.generate_markdown_var.get():
            # Markdown enabled: enable simple markdown checkbox
            self.simple_markdown_checkbox.config(state=tk.NORMAL)
            
            if self.simple_markdown_var.get():
                # Simple markdown mode: disable GPU and models
                self.use_gpu_checkbox.config(state=tk.DISABLED)
                self.preload_models_btn.config(state=tk.DISABLED)
                if log_message:
                    self.print_to_console("[INFO] Simple Markdown mode: Using PyMuPDF4LLM (no OCR, GPU not needed)", "info")
            else:
                # Advanced markdown mode: enable GPU and models
                self.use_gpu_checkbox.config(state=tk.NORMAL)
                self.preload_models_btn.config(state=tk.NORMAL)
        else:
            # Markdown disabled: gray out everything
            self.simple_markdown_checkbox.config(state=tk.DISABLED)
            self.use_gpu_checkbox.config(state=tk.DISABLED)
            self.preload_models_btn.config(state=tk.DISABLED)

    # New: Method to handle GPU checkbox changes
    def on_gpu_checkbox_change(self):
        """Handles changes to the 'Use GPU' checkbox state."""
        import torch
        if self.use_gpu_var.get() and not torch.cuda.is_available():
            self.print_to_console("[WARNING] GPU acceleration requested but CUDA not available. Will use CPU.", "warning")
        elif self.use_gpu_var.get():
            self.print_to_console(f"[INFO] GPU acceleration enabled. Using device: cuda:0", "info")
        else:
            self.print_to_console("[INFO] Using CPU for processing.", "info")
        
        self.log_and_save_setting("GPU Acceleration", self.use_gpu_var)
        
        # Clear existing models to force reload with new device settings
        if hasattr(self, '_marker_models'):
            delattr(self, '_marker_models')
            if hasattr(self, '_device'):
                delattr(self, '_device')
            # Update UI based on whether models exist on disk
            if self._check_models_exist():
                self._update_models_ui_found()
            else:
                self._update_models_ui_not_found()

    def _check_models_exist(self):
        """Check if marker-pdf models exist in the selected directory."""
        try:
            if not os.path.exists(self.models_directory):
                return False
            
            # Check for actual model files/directories that Surya creates
            # Look for the specific directory structure that Surya uses
            expected_patterns = [
                'models--datalab-to--surya-layout',
                'models--datalab-to--surya-det',
                'models--datalab-to--surya-rec',
                'models--datalab-to--surya-table-rec',
                'models--datalab-to--surya-ocr-error',
                # Alternative patterns
                'layout',
                'detection', 
                'recognition',
                'table_rec',
                'ocr_error'
            ]
            
            found_models = []
            for item in os.listdir(self.models_directory):
                item_path = os.path.join(self.models_directory, item)
                if os.path.isdir(item_path):
                    for pattern in expected_patterns:
                        if pattern in item.lower():
                            found_models.append(item)
                            break
            
            # If we found at least 3 model directories, consider models present
            models_present = len(found_models) >= 3
            
            if models_present:
                self.print_to_console(f"[INFO] Found {len(found_models)} model directories in {self.models_directory}", "info")
            
            return models_present
            
        except (OSError, FileNotFoundError):
            return False

    def _update_models_ui_found(self):
        """Update UI when models are found or loaded."""
        self.preload_models_btn.config(text="Change Folder", state=tk.NORMAL)
        # Show models path in green
        path_text = f"✓ Models found: {self.models_directory}"
        self.models_path_label.config(text=path_text, fg="green")
        
    def _update_models_ui_not_found(self):
        """Update UI when models are not found."""
        self.preload_models_btn.config(text="Select Models Dir", state=tk.NORMAL)
        self.models_path_label.config(text="", fg="green")

    def preload_marker_models(self):
        """Preloads marker-pdf models with folder selection."""
        # Show folder picker for models directory
        selected_folder = filedialog.askdirectory(
            title="Select Models Directory",
            initialdir=self.models_directory
        )
        
        if not selected_folder:
            return  # User cancelled
        
        # Update models directory
        self.models_directory = selected_folder
        self.save_settings()  # Save the selection
        
        self.print_to_console(f"[INFO] Selected models directory: {self.models_directory}", "info")
        
        # Check if models already exist
        if self._check_models_exist():
            self.print_to_console("[INFO] Models already exist in selected directory.", "success")
            self._update_models_ui_found()
            return
        
        # Models don't exist, start download
        if hasattr(self, '_marker_models') and self._marker_models:
            self.print_to_console("[INFO] Marker-pdf models already loaded in memory.", "info")
            return
        
        self.print_to_console("[INFO] Starting marker-pdf model download...", "info")
        self.preload_models_btn.config(state=tk.DISABLED, text="Downloading...")
        
        def preload_thread():
            try:
                # CRITICAL: Set environment variables BEFORE any marker imports
                os.makedirs(self.models_directory, exist_ok=True)
                self.master.after(0, lambda: self.print_to_console(f"[INFO] Using models directory: {self.models_directory}", "info"))
                
                # Set environment variables to use selected models directory
                os.environ['TORCH_HOME'] = self.models_directory
                os.environ['HF_HOME'] = self.models_directory
                os.environ['TRANSFORMERS_CACHE'] = self.models_directory
                # IMPORTANT: Set Surya model cache directory
                os.environ['MODEL_CACHE_DIR'] = self.models_directory
                
                # CRITICAL: Disable multiprocessing to prevent process pool errors in frozen executable
                os.environ['MARKER_NO_MULTIPROCESSING'] = '1'
                os.environ['OMP_NUM_THREADS'] = '1'
                os.environ['MKL_NUM_THREADS'] = '1'
                
                self.master.after(0, lambda: self.print_to_console(f"[INFO] Set MODEL_CACHE_DIR to: {self.models_directory}", "info"))
                
                # Verify environment variable is set
                actual_cache_dir = os.environ.get('MODEL_CACHE_DIR')
                self.master.after(0, lambda: self.print_to_console(f"[DEBUG] MODEL_CACHE_DIR env var: {actual_cache_dir}", "debug"))
                
                # Now import torch and determine device
                import torch
                if self.use_gpu_var.get() and torch.cuda.is_available():
                    device = "cuda"
                    os.environ['TORCH_DEVICE'] = 'cuda'
                    self.master.after(0, lambda: self.print_to_console(f"[INFO] Using GPU acceleration: {torch.cuda.get_device_name(0)}", "info"))
                else:
                    device = "cpu"
                    os.environ['TORCH_DEVICE'] = 'cpu'
                    if self.use_gpu_var.get():
                        self.master.after(0, lambda: self.print_to_console("[WARNING] GPU requested but not available, using CPU", "warning"))
                    else:
                        self.master.after(0, lambda: self.print_to_console("[INFO] Using CPU for processing", "info"))
                
                # Set up output capture for download progress
                import sys
                from contextlib import redirect_stdout, redirect_stderr
                
                # Create custom output handlers for the preload thread
                class ThreadSafePreloadCapture:
                    def __init__(self, console_func, master, tag="progress"):
                        self.console_func = console_func
                        self.master = master
                        self.tag = tag
                    
                    def write(self, text):
                        if text.strip():
                            # Schedule GUI update in main thread
                            self.master.after(0, lambda: self.console_func(f"[DOWNLOAD] {text.strip()}", self.tag))
                    
                    def flush(self):
                        pass
                
                stdout_capture = ThreadSafePreloadCapture(self.print_to_console, self.master, "progress")
                stderr_capture = ThreadSafePreloadCapture(self.print_to_console, self.master, "warning")
                
                # NOW import marker modules after environment is set
                self.master.after(0, lambda: self.print_to_console("[INFO] Importing marker modules with new environment...", "progress"))
                
                # Import marker modules and do all model operations with output capture
                with redirect_stdout(stdout_capture), redirect_stderr(stderr_capture):
                    from marker.models import create_model_dict
                    
                    # Check where Surya thinks it should store models
                    from surya import settings
                    actual_model_dir = settings.settings.MODEL_CACHE_DIR
                    self.master.after(0, lambda: self.print_to_console(f"[DEBUG] Surya MODEL_CACHE_DIR: {actual_model_dir}", "debug"))
                    
                    self.master.after(0, lambda: self.print_to_console("[INFO] Downloading and initializing marker-pdf models...", "progress"))
                    
                    # Create model dictionary and force download by actually using the models
                    from marker.converters.pdf import PdfConverter
                    from marker.output import text_from_rendered
                    
                    models = create_model_dict()
                    self.master.after(0, lambda: self.print_to_console("[INFO] Models created, testing with sample conversion...", "progress"))
                    
                    # Create converter to trigger actual model downloads
                    converter = PdfConverter(artifact_dict=models)
                    
                    # Test with a minimal PDF to ensure models are fully downloaded
                    # Create a simple test PDF in memory
                    import tempfile
                    test_pdf_content = b"""%PDF-1.4
1 0 obj
<<
/Type /Catalog
/Pages 2 0 R
>>
endobj
2 0 obj
<<
/Type /Pages
/Kids [3 0 R]
/Count 1
>>
endobj
3 0 obj
<<
/Type /Page
/Parent 2 0 R
/MediaBox [0 0 612 792]
/Contents 4 0 R
>>
endobj
4 0 obj
<<
/Length 44
>>
stream
BT
/F1 12 Tf
100 700 Td
(Test) Tj
ET
endstream
endobj
xref
0 5
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000206 00000 n 
trailer
<<
/Size 5
/Root 1 0 R
>>
startxref
300
%%EOF"""
                    
                    with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as temp_pdf:
                        temp_pdf.write(test_pdf_content)
                        temp_pdf_path = temp_pdf.name
                    
                    try:
                        self.master.after(0, lambda: self.print_to_console("[INFO] Running test conversion to download models...", "progress"))
                        rendered = converter(temp_pdf_path)
                        text, _, images = text_from_rendered(rendered)
                        self.master.after(0, lambda: self.print_to_console("[OK] Test conversion successful - models fully downloaded!", "success"))
                        
                        # Store the models for later use
                        self._marker_models = models
                        self._device = device
                        
                    finally:
                        # Clean up test file
                        try:
                            os.unlink(temp_pdf_path)
                        except:
                            pass
                self.master.after(0, lambda: self._on_preload_complete(True))
                
            except Exception as e:
                self.master.after(0, lambda: self._on_preload_complete(False, str(e)))
        
        threading.Thread(target=preload_thread, daemon=True).start()

    def _on_preload_complete(self, success, error_msg=None):
        """Called when model preloading completes."""
        if success:
            self.print_to_console("[OK] Marker-pdf models downloaded successfully!", "success")
            self._update_models_ui_found()
        else:
            self.print_to_console(f"[ERROR] Failed to download marker-pdf models: {error_msg}", "error")
            self._update_models_ui_not_found()

    def on_listbox_select(self, event):
        """Enables/disables buttons based on listbox selection."""
        selected_indices = self.pdf_listbox.curselection()
        has_selection = len(selected_indices) > 0
        list_size = self.pdf_listbox.size()
        
        if has_selection:
            selected_idx = selected_indices[0]
            # Enable move buttons
            self.move_up_btn.config(state=tk.NORMAL if selected_idx > 0 else tk.DISABLED)
            self.move_down_btn.config(state=tk.NORMAL if selected_idx < list_size - 1 else tk.DISABLED)
            self.move_to_top_btn.config(state=tk.NORMAL if selected_idx > 0 else tk.DISABLED)
            self.move_to_bottom_btn.config(state=tk.NORMAL if selected_idx < list_size - 1 else tk.DISABLED)
            self.remove_btn.config(state=tk.NORMAL)
        else:
            self.move_up_btn.config(state=tk.DISABLED)
            self.move_down_btn.config(state=tk.DISABLED)
            self.move_to_top_btn.config(state=tk.DISABLED)
            self.move_to_bottom_btn.config(state=tk.DISABLED)
            self.remove_btn.config(state=tk.DISABLED)
        
        self.clear_all_btn.config(state=tk.NORMAL if list_size > 0 else tk.DISABLED)

    def _get_log_level_from_tag(self, tag):
        """Maps console tags to standard logging levels."""
        tag_lower = (tag or "").lower()
        level_mapping = {
            "debug": "DEBUG",
            "info": "INFO",
            "warning": "WARNING",
            "error": "ERROR",
            "success": "INFO",  # Success messages are typically INFO level
            "progress": "INFO",  # Progress messages are typically INFO level
            None: "INFO",  # Default to INFO if no tag
        }
        return level_mapping.get(tag_lower, "INFO")
    
    def _should_show_message(self, level):
        """Determines if a message should be shown based on current filter."""
        filter_level = self.console_filter_level_var.get()
        if filter_level == "ALL":
            return True
        
        # Define level hierarchy (lower number = higher priority)
        level_hierarchy = {
            "DEBUG": 0,
            "INFO": 1,
            "WARNING": 2,
            "ERROR": 3,
            "CRITICAL": 4
        }
        
        filter_value = level_hierarchy.get(filter_level, 1)
        message_value = level_hierarchy.get(level, 1)
        
        # Show messages at or above the filter level
        return message_value >= filter_value
    
    def print_to_console(self, message, tag=None):
        """Prints a message to the console output widget with filtering support."""
        # Determine log level from tag
        level = self._get_log_level_from_tag(tag)
        
        # Store message in buffer
        self.console_message_buffer.append((message, tag, level))
        
        # Check if message should be shown based on filter
        if not self._should_show_message(level):
            return
        
        # Only insert if console is visible
        if self.console_visible_var.get():
            self.console_output.insert(tk.END, message + "\n", tag)
            self.console_output.see(tk.END)
    
    def _toggle_console_visibility(self):
        """Shows or hides the console based on checkbox state."""
        if self.console_visible_var.get():
            self.console_output.pack(fill=tk.BOTH, expand=True)
            self._refresh_console_display()
        else:
            self.console_output.pack_forget()
        self.save_settings()
    
    def _on_filter_change(self, *args):
        """Called when filter level changes - refreshes the console display."""
        self._refresh_console_display()
        self.save_settings()
    
    def _refresh_console_display(self):
        """Refreshes the console display based on current filter and visibility."""
        if not self.console_visible_var.get():
            return
        
        # Clear current display
        self.console_output.delete(1.0, tk.END)
        
        # Re-insert messages that match the current filter
        filter_level = self.console_filter_level_var.get()
        for message, tag, level in self.console_message_buffer:
            if self._should_show_message(level):
                self.console_output.insert(tk.END, message + "\n", tag)
        
        # Scroll to end
        self.console_output.see(tk.END)

    def log_and_save_setting(self, setting_name, var):
        """Logs checkbox state changes and saves all settings."""
        state = "Enabled" if var.get() else "Disabled"
        self.print_to_console(f"Configuration: {setting_name} {state}.", "info")
        self.save_settings()

    def load_settings(self):
        """Loads settings from settings.json."""
        if os.path.exists(SETTINGS_FILE):
            try:
                with open(SETTINGS_FILE, "r") as f:
                    settings = json.load(f)
                    self.pdf_files = settings.get("pdf_files", [])
                    self.input_folder = settings.get("input_folder", DOWNLOADS_PATH)
                    self.output_folder = settings.get("output_folder", DOWNLOADS_PATH)
                    self.remove_timestamps_var.set(settings.get("remove_timestamps_enabled", False))
                    self.remove_images_var.set(settings.get("remove_images_enabled", False))
                    self.remove_pii_var.set(settings.get("remove_pii_enabled", False))
                    self.custom_pii_var.set(settings.get("custom_pii_strings", ""))
                    # New: Load split settings
                    self.split_by_words_var.set(settings.get("split_by_words_enabled", False))
                    self.split_word_count_var.set(settings.get("split_word_count", "10000"))
                    # New: Load markdown setting
                    self.generate_markdown_var.set(settings.get("generate_markdown_enabled", False))
                    # New: Load simple markdown setting
                    self.simple_markdown_var.set(settings.get("simple_markdown_enabled", False))
                    # New: Load markdown type
                    self.markdown_type_var.set(settings.get("markdown_type", "simple"))
                    # New: Load GPU setting
                    self.use_gpu_var.set(settings.get("use_gpu_enabled", False))
                    # New: Load models directory
                    self.models_directory = settings.get("models_directory", MODELS_DIR)
                    # New: Load qpdf path
                    self.qpdf_path = settings.get("qpdf_path", None)
                    # New: Load console settings
                    self.console_visible_var.set(settings.get("console_visible", True))
                    self.console_filter_level_var.set(settings.get("console_filter_level", "ALL"))
                    # New: Load multi-format settings
                    self.output_file_type_var.set(settings.get("output_file_type", "PDF"))
                    self.output_filename_var.set(settings.get("output_filename", ""))
                    self.preserve_formatting_var.set(settings.get("preserve_formatting", False))

                    self.print_to_console(f"Loaded settings from {SETTINGS_FILE}", "info")

                    self.total_word_count = 0
                    files_to_keep = []
                    self.pdf_listbox.delete(0, tk.END)
                    for file_path in self.pdf_files:
                        if os.path.exists(file_path):
                            try:
                                text = self._extract_text_from_file(file_path)
                                self.total_word_count += self._count_words(text)
                                files_to_keep.append(file_path)
                                self.pdf_listbox.insert(tk.END, os.path.basename(file_path))
                            except Exception as e:
                                self.print_to_console(f"Error processing {os.path.basename(file_path)} on load: {e}", "error")
                        else:
                            self.print_to_console(f"Warning: Stored file not found: {file_path}. Removing from list.", "warning")
                    self.pdf_files = files_to_keep
            except Exception as e:
                self.print_to_console(f"Error loading settings: {e}. Starting with defaults.", "error")
                self.pdf_files = []
                self.input_folder = DOWNLOADS_PATH
                self.output_folder = DOWNLOADS_PATH

        self.input_folder_label.config(text=self.input_folder)
        self.output_folder_label.config(text=self.output_folder)
        self.on_listbox_select(None)
        
        # Check if models exist on startup
        if self._check_models_exist():
            self._update_models_ui_found()
        else:
            self._update_models_ui_not_found()
        
        # Update markdown controls visibility based on loaded settings
        self._update_markdown_controls_visibility()
        
        # Update qpdf UI status after loading settings
        self._update_qpdf_ui_status()
        
        # Update console visibility and refresh display after loading settings
        self._toggle_console_visibility()
        self._refresh_console_display()

    def save_settings(self):
        """Saves current settings to settings.json."""
        settings = {
            "pdf_files": self.pdf_files,
            "input_folder": self.input_folder,
            "output_folder": self.output_folder,
            "remove_timestamps_enabled": self.remove_timestamps_var.get(),
            "remove_images_enabled": self.remove_images_var.get(),
            "remove_pii_enabled": self.remove_pii_var.get(),
            "custom_pii_strings": self.custom_pii_var.get(),
            # New: Save split settings
            "split_by_words_enabled": self.split_by_words_var.get(),
            "split_word_count": self.split_word_count_var.get(),
            # New: Save markdown setting
            "generate_markdown_enabled": self.generate_markdown_var.get(),
            # New: Save simple markdown setting
            "simple_markdown_enabled": self.simple_markdown_var.get(),
            # New: Save markdown type
            "markdown_type": self.markdown_type_var.get(),
            # New: Save GPU setting
            "use_gpu_enabled": self.use_gpu_var.get(),
            # New: Save models directory
            "models_directory": self.models_directory,
            # New: Save qpdf path
            "qpdf_path": self.qpdf_path,
            # New: Save console settings
            "console_visible": self.console_visible_var.get(),
            "console_filter_level": self.console_filter_level_var.get(),
            # New: Save multi-format settings
            "output_file_type": self.output_file_type_var.get(),
            "output_filename": self.output_filename_var.get(),
            "preserve_formatting": self.preserve_formatting_var.get()
        }
        try:
            with open(SETTINGS_FILE, "w") as f:
                json.dump(settings, f, indent=4)
        except Exception as e:
            self.print_to_console(f"Error saving settings: {e}", "error")

    def _extract_text_from_file(self, file_path):
        """Extracts text from any supported file format."""
        ext = os.path.splitext(file_path)[1].lower()

        if ext == '.pdf':
            return self._extract_text_from_pdf(file_path)
        elif ext == '.txt':
            return self._extract_text_from_txt(file_path)
        elif ext == '.md':
            return self._extract_text_from_md(file_path)
        elif ext == '.docx':
            return self._extract_text_from_docx(file_path)
        elif ext == '.odt':
            return self._extract_text_from_odt(file_path)
        elif ext == '.rtf':
            return self._extract_text_from_rtf(file_path)
        elif ext == '.epub':
            return self._extract_text_from_epub(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _extract_text_from_txt(self, file_path):
        """Extracts text from TXT file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_text_from_md(self, file_path):
        """Extracts text from Markdown file."""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_text_from_docx(self, file_path):
        """Extracts text from DOCX file."""
        try:
            from docx import Document
            doc = Document(file_path)
            text = []
            for paragraph in doc.paragraphs:
                text.append(paragraph.text)
            return '\n'.join(text)
        except Exception as e:
            raise Exception(f"Error reading DOCX file: {e}")

    def _extract_text_from_odt(self, file_path):
        """Extracts text from ODT file."""
        try:
            from odf import text, teletype
            from odf.opendocument import load
            doc = load(file_path)
            all_text = []
            for paragraph in doc.getElementsByType(text.P):
                all_text.append(teletype.extractText(paragraph))
            return '\n'.join(all_text)
        except Exception as e:
            raise Exception(f"Error reading ODT file: {e}")

    def _extract_text_from_rtf(self, file_path):
        """Extracts text from RTF file."""
        try:
            from striprtf.striprtf import rtf_to_text
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            return rtf_to_text(rtf_content)
        except Exception as e:
            raise Exception(f"Error reading RTF file: {e}")

    def _extract_text_from_epub(self, file_path):
        """Extracts text from EPUB file."""
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup

            book = epub.read_epub(file_path)
            text_content = []

            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text_content.append(soup.get_text())

            return '\n'.join(text_content)
        except Exception as e:
            raise Exception(f"Error reading EPUB file: {e}")

    def _extract_text_from_pdf(self, pdf_path):
        """Extracts text from a PDF for word counting."""
        text = ""
        timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
        try:
            # Modified: Can now accept a document object or a path
            if isinstance(pdf_path, str):
                doc = fitz.open(pdf_path)
                close_doc = True
            else:
                doc = pdf_path
                close_doc = False

            for page in doc:
                page_text = page.get_text("text")
                if self.remove_timestamps_var.get():
                    page_text = re.sub(timestamp_regex, '', page_text)
                text += page_text + " "

            if close_doc:
                doc.close()
        except Exception as e:
            self.print_to_console(f"Error extracting text from PDF: {e}", "error")
            if 'close_doc' in locals() and close_doc and 'doc' in locals():
                doc.close()
            raise
        return text

    def _count_words(self, text):
        """Counts words in a given text string."""
        return len(re.findall(r'\b\w+\b', text.lower()))

    def _generate_output_file(self, text, output_filepath):
        """Generate output file in the selected format."""
        output_type = self.output_file_type_var.get().lower()

        if output_type == 'pdf':
            self._generate_pdf(text, output_filepath)
        elif output_type == 'txt':
            self._generate_txt(text, output_filepath)
        elif output_type == 'md':
            self._generate_md(text, output_filepath)
        elif output_type == 'docx':
            self._generate_docx(text, output_filepath)
        elif output_type == 'odt':
            self._generate_odt(text, output_filepath)
        elif output_type == 'rtf':
            self._generate_rtf(text, output_filepath)
        elif output_type == 'epub':
            self._generate_epub(text, output_filepath)
        else:
            raise ValueError(f"Unsupported output format: {output_type}")

    def _generate_pdf(self, text, output_filepath):
        """Generate PDF output from text."""
        doc = fitz.open()
        page = doc.new_page()
        text_rect = fitz.Rect(50, 50, page.rect.width - 50, page.rect.height - 50)
        page.insert_textbox(text_rect, text, fontsize=11, fontname="helv")
        doc.save(output_filepath)
        doc.close()

    def _generate_txt(self, text, output_filepath):
        """Generate TXT output from text."""
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write(text)

    def _generate_md(self, text, output_filepath):
        """Generate MD (Markdown) output from text."""
        with open(output_filepath, 'w', encoding='utf-8') as f:
            f.write(text)

    def _generate_docx(self, text, output_filepath):
        """Generate DOCX output from text."""
        from docx import Document
        doc = Document()
        # Split text by paragraphs and add to document
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para)
        doc.save(output_filepath)

    def _generate_odt(self, text, output_filepath):
        """Generate ODT output from text."""
        from odf.opendocument import OpenDocumentText
        from odf.text import P
        from odf import text as odf_text

        doc = OpenDocumentText()
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                p = P(text=para)
                doc.text.appendChild(p)
        doc.save(output_filepath)

    def _generate_rtf(self, text, output_filepath):
        """Generate RTF output from text using pypandoc."""
        try:
            import pypandoc
            # Write text to a temporary file
            temp_txt = output_filepath + ".tmp.txt"
            with open(temp_txt, 'w', encoding='utf-8') as f:
                f.write(text)
            # Convert using pypandoc
            pypandoc.convert_file(temp_txt, 'rtf', outputfile=output_filepath)
            # Clean up temp file
            os.remove(temp_txt)
        except Exception as e:
            # Fallback to basic RTF if pypandoc fails
            with open(output_filepath, 'w', encoding='utf-8') as f:
                f.write(r'{\rtf1\ansi\deff0 {\fonttbl {\f0 Times New Roman;}}')
                f.write(r'\f0\fs24 ')
                # Escape special RTF characters
                rtf_text = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                f.write(rtf_text)
                f.write(r'}')

    def _generate_epub(self, text, output_filepath):
        """Generate EPUB output from text."""
        import ebooklib
        from ebooklib import epub

        book = epub.EpubBook()
        book.set_identifier('merged_document')
        book.set_title('Merged Document')
        book.set_language('en')

        # Create chapter
        c1 = epub.EpubHtml(title='Chapter 1', file_name='chap_01.xhtml', lang='en')
        # Convert text to HTML paragraphs
        html_content = '<h1>Merged Document</h1>'
        paragraphs = text.split('\n')
        for para in paragraphs:
            if para.strip():
                html_content += f'<p>{para}</p>'
        c1.content = html_content

        # Add chapter to book
        book.add_item(c1)
        book.toc = (epub.Link('chap_01.xhtml', 'Chapter 1', 'chap_01'),)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ['nav', c1]

        # Write EPUB file
        epub.write_epub(output_filepath, book, {})

    def update_word_count_display(self):
        """Updates the total word count label in the GUI."""
        self.total_words_label.config(text=f"Total Words: {self.total_word_count}")

    def add_pdf_file(self):
        """Adds selected files (PDF, ODT, DOCX, TXT, RTF, EPUB, MD) to the list."""
        filetypes = [
            ("All Supported Files", "*.pdf *.odt *.docx *.txt *.rtf *.epub *.md"),
            ("PDF files", "*.pdf"),
            ("ODT files", "*.odt"),
            ("Word files", "*.docx"),
            ("Text files", "*.txt"),
            ("RTF files", "*.rtf"),
            ("EPUB files", "*.epub"),
            ("Markdown files", "*.md"),
            ("All files", "*.*")
        ]
        file_paths = filedialog.askopenfilenames(
            title="Select Files",
            filetypes=filetypes,
            initialdir=self.input_folder
        )
        if not file_paths:
            return

        for file_path in file_paths:
            if file_path not in self.pdf_files:
                self.pdf_files.append(file_path)
                self.pdf_listbox.insert(tk.END, os.path.basename(file_path))
                self.print_to_console(f"Added: {os.path.basename(file_path)}", "info")
                try:
                    text = self._extract_text_from_file(file_path)
                    words_in_file = self._count_words(text)
                    self.total_word_count += words_in_file
                    self.print_to_console(f"  - Words in '{os.path.basename(file_path)}': {words_in_file}", "info")
                except Exception as e:
                    self.print_to_console(f"Could not count words for {os.path.basename(file_path)}: {e}", "error")
                    self.pdf_files.remove(file_path)
                    self.pdf_listbox.delete(tk.END)
            else:
                self.print_to_console(f"'{os.path.basename(file_path)}' is already in the list.", "info")

        self.update_word_count_display()
        self.save_settings()
        self.on_listbox_select(None)

    def remove_pdf_file(self):
        """Removes selected PDF files from the list."""
        selected_indices = sorted(self.pdf_listbox.curselection(), reverse=True)
        if not selected_indices:
            return

        for index in selected_indices:
            removed_path = self.pdf_files.pop(index)
            self.pdf_listbox.delete(index)
            self.print_to_console(f"Removed: {os.path.basename(removed_path)}", "info")
        
        self.total_word_count = 0
        for pdf_path in self.pdf_files:
            try:
                text = self._extract_text_from_pdf(pdf_path) 
                self.total_word_count += self._count_words(text)
            except Exception as e:
                self.print_to_console(f"Error recalculating words for {os.path.basename(pdf_path)}: {e}", "error")

        self.update_word_count_display()
        self.save_settings()
        self.on_listbox_select(None)

    def clear_all_pdfs(self):
        """Clears all PDF files from the list."""
        if messagebox.askyesno("Clear All", "Are you sure you want to remove all PDF files?"):
            self.pdf_files.clear()
            self.pdf_listbox.delete(0, tk.END)
            self.total_word_count = 0
            self.update_word_count_display()
            self.save_settings()
            self.print_to_console("All PDF files cleared.", "info")
            self.on_listbox_select(None)

    def move_pdf_up(self):
        """Moves selected item up."""
        selected_indices = self.pdf_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index > 0:
            self.pdf_files[index], self.pdf_files[index-1] = self.pdf_files[index-1], self.pdf_files[index]
            self.pdf_listbox.delete(index)
            self.pdf_listbox.insert(index-1, os.path.basename(self.pdf_files[index-1]))
            self.pdf_listbox.selection_set(index-1)
            self.save_settings()

    def move_pdf_down(self):
        """Moves selected item down."""
        selected_indices = self.pdf_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index < len(self.pdf_files) - 1:
            self.pdf_files[index], self.pdf_files[index+1] = self.pdf_files[index+1], self.pdf_files[index]
            self.pdf_listbox.delete(index)
            self.pdf_listbox.insert(index+1, os.path.basename(self.pdf_files[index+1]))
            self.pdf_listbox.selection_set(index+1)
            self.save_settings()
    
    def move_pdf_to_top(self):
        """Moves selected item to the top of the list."""
        selected_indices = self.pdf_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index > 0:
            # Remove item from current position
            pdf_path = self.pdf_files.pop(index)
            self.pdf_listbox.delete(index)
            # Insert at top
            self.pdf_files.insert(0, pdf_path)
            self.pdf_listbox.insert(0, os.path.basename(pdf_path))
            self.pdf_listbox.selection_set(0)
            self.save_settings()
    
    def move_pdf_to_bottom(self):
        """Moves selected item to the bottom of the list."""
        selected_indices = self.pdf_listbox.curselection()
        if not selected_indices: return
        index = selected_indices[0]
        if index < len(self.pdf_files) - 1:
            # Remove item from current position
            pdf_path = self.pdf_files.pop(index)
            self.pdf_listbox.delete(index)
            # Insert at bottom
            self.pdf_files.append(pdf_path)
            self.pdf_listbox.insert(tk.END, os.path.basename(pdf_path))
            self.pdf_listbox.selection_set(len(self.pdf_files) - 1)
            self.save_settings()

    def select_output_folder(self):
        """Opens a dialog to select the output folder."""
        folder_selected = filedialog.askdirectory(initialdir=self.output_folder)
        if folder_selected:
            self.output_folder = folder_selected
            self.output_folder_label.config(text=self.output_folder)
            self.print_to_console(f"Output folder set to: {self.output_folder}", "info")
            self.save_settings()
    
    def _check_qpdf_executable(self, qpdf_path):
        """Validates that a qpdf executable works by testing it."""
        if not qpdf_path or not os.path.exists(qpdf_path):
            return False
        
        try:
            # Test if qpdf works by calling it with --version
            result = subprocess.run(
                [qpdf_path, "--version"],
                capture_output=True,
                text=True,
                timeout=5
            )
            return result.returncode == 0
        except (subprocess.TimeoutExpired, FileNotFoundError, OSError):
            return False
    
    def _update_qpdf_ui_status(self):
        """Updates the UI to show the current qpdf status and updates button accordingly."""
        if self.qpdf_path and self._check_qpdf_executable(self.qpdf_path):
            # Show shortened path if too long
            display_path = self.qpdf_path
            if len(display_path) > 50:
                display_path = "..." + display_path[-47:]
            self.qpdf_path_label.config(text=f"✓ {os.path.basename(display_path)}", fg="green")
            # Change button to "Decrypt PDF"
            self.decrypt_btn.config(text="Decrypt PDF", command=self.decrypt_pdf)
            self.print_to_console(f"[INFO] qpdf configured: {self.qpdf_path}", "info")
        else:
            self.qpdf_path_label.config(text="Not configured", fg="red")
            # Change button to "Locate qpdf"
            self.decrypt_btn.config(text="Locate qpdf", command=self._decrypt_or_locate)
            if not self.qpdf_path:
                self.print_to_console("[INFO] qpdf not configured. Click 'Locate qpdf' to set it up.", "warning")
    
    def _decrypt_or_locate(self):
        """Routes to locate qpdf if not configured, otherwise decrypt."""
        if not self.qpdf_path or not self._check_qpdf_executable(self.qpdf_path):
            self.locate_qpdf_executable()
        else:
            self.decrypt_pdf()
    
    def locate_qpdf_executable(self):
        """Opens a file dialog to locate and configure the qpdf executable."""
        # Determine initial directory
        initial_dir = None
        if self.qpdf_path:
            initial_dir = os.path.dirname(self.qpdf_path)
        elif os.path.exists(os.path.join(os.environ.get("ProgramFiles", ""), "qpdf")):
            initial_dir = os.path.join(os.environ.get("ProgramFiles", ""), "qpdf")
        else:
            initial_dir = os.path.expanduser("~")
        
        # Open file dialog to select qpdf executable
        qpdf_path = filedialog.askopenfilename(
            title="Locate qpdf Executable",
            initialdir=initial_dir,
            filetypes=[
                ("Executable files", "*.exe"),
                ("All files", "*.*")
            ]
        )
        
        if not qpdf_path:
            return  # User cancelled
        
        # Validate the selected executable
        if not self._check_qpdf_executable(qpdf_path):
            messagebox.showerror(
                "Invalid qpdf Executable",
                f"The selected file does not appear to be a valid qpdf executable.\n\n"
                f"Selected: {qpdf_path}\n\n"
                f"Please ensure you select the qpdf.exe file from your qpdf installation."
            )
            self.print_to_console(f"[ERROR] Invalid qpdf executable selected: {qpdf_path}", "error")
            return
        
        # Save the path
        self.qpdf_path = qpdf_path
        self.save_settings()
        self._update_qpdf_ui_status()
        self.print_to_console(f"[SUCCESS] qpdf executable configured: {qpdf_path}", "success")
    
    def _open_qpdf_download_page(self):
        """Opens the qpdf releases page in the default browser."""
        url = "https://github.com/qpdf/qpdf/releases/"
        try:
            webbrowser.open(url)
            self.print_to_console(f"[INFO] Opening qpdf download page: {url}", "info")
        except Exception as e:
            self.print_to_console(f"[ERROR] Could not open browser: {e}", "error")
            messagebox.showerror(
                "Error",
                f"Could not open browser. Please manually visit:\n\n{url}\n\nError: {e}"
            )
    
    def decrypt_pdf(self):
        """Opens a file dialog to select a PDF file, decrypts it with qpdf, and saves to output folder."""
        # Check if qpdf is configured and available
        if not self.qpdf_path or not self._check_qpdf_executable(self.qpdf_path):
            messagebox.showerror(
                "qpdf Not Configured",
                "qpdf executable is not configured.\n\n"
                "Please click 'Locate qpdf executable' to select the qpdf.exe file.\n\n"
                "You can download qpdf from the releases page (click 'Get qpdf for Decryption')."
            )
            self.print_to_console("[ERROR] qpdf not configured. Please locate the qpdf executable first.", "error")
            return
        
        qpdf_path = self.qpdf_path
        
        # Open file dialog to select PDF
        pdf_path = filedialog.askopenfilename(
            title="Select PDF File to Decrypt",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        if not pdf_path:
            return  # User cancelled
        
        # Check if file is encrypted
        is_encrypted = False
        try:
            doc = fitz.open(pdf_path)
            is_encrypted = doc.needs_pass
            doc.close()
        except Exception as e:
            # If we can't open it, it might be encrypted
            error_str = str(e).lower()
            if "password" in error_str or "encrypted" in error_str:
                is_encrypted = True
            else:
                self.print_to_console(f"[ERROR] Could not read PDF file: {e}", "error")
                messagebox.showerror("Error", f"Could not read PDF file: {e}")
                return
        
        if not is_encrypted:
            response = messagebox.askyesno(
                "PDF Not Encrypted",
                "This PDF file does not appear to be encrypted.\n\n"
                "Do you still want to process it with qpdf? (This will still create a decrypted copy)"
            )
            if not response:
                return
        
        # Prompt for password if needed
        password = None
        if is_encrypted:
            password = simpledialog.askstring(
                "PDF Password",
                "This PDF is encrypted. Enter the password (leave empty for no password):",
                show="*"
            )
            if password is None:
                return  # User cancelled
            # Convert empty string to None for qpdf (empty string means no password attempt)
            if password == "":
                password = None
        
        # Generate output filename
        base_name = os.path.splitext(os.path.basename(pdf_path))[0]
        output_filename = f"{base_name}_decrypted.pdf"
        output_path = os.path.join(self.output_folder, output_filename)
        
        # Ensure unique filename
        counter = 1
        while os.path.exists(output_path):
            output_filename = f"{base_name}_decrypted_{counter}.pdf"
            output_path = os.path.join(self.output_folder, output_filename)
            counter += 1
        
        # Run decryption in a separate thread to avoid blocking UI
        self.print_to_console(f"[INFO] Starting PDF decryption: {os.path.basename(pdf_path)}", "info")
        
        def decrypt_thread():
            try:
                # Build qpdf command
                cmd = [qpdf_path]
                
                # Add password if provided (qpdf requires --password=password format)
                if password:
                    cmd.append(f"--password={password}")
                
                # Add decrypt flag and files
                cmd.extend(["--decrypt", pdf_path, output_path])
                
                self.master.after(0, lambda: self.print_to_console(f"[PROGRESS] Running qpdf decryption...", "progress"))
                
                # Run qpdf
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=60  # 60 second timeout
                )
                
                if result.returncode == 0:
                    self.master.after(0, lambda: self.print_to_console(
                        f"[SUCCESS] PDF decrypted successfully: {output_filename}", "success"
                    ))
                    self.master.after(0, lambda: messagebox.showinfo(
                        "Success",
                        f"PDF decrypted successfully!\n\nSaved to:\n{output_path}"
                    ))
                else:
                    error_msg = result.stderr.strip() if result.stderr else result.stdout.strip()
                    self.master.after(0, lambda: self.print_to_console(
                        f"[ERROR] qpdf decryption failed: {error_msg}", "error"
                    ))
                    self.master.after(0, lambda: messagebox.showerror(
                        "Decryption Failed",
                        f"Failed to decrypt PDF:\n\n{error_msg}\n\n"
                        "Please check if the password is correct or if the file is encrypted."
                    ))
                    
            except subprocess.TimeoutExpired:
                self.master.after(0, lambda: self.print_to_console(
                    "[ERROR] qpdf operation timed out", "error"
                ))
                self.master.after(0, lambda: messagebox.showerror(
                    "Timeout",
                    "The decryption operation timed out. The file may be too large or corrupted."
                ))
            except Exception as e:
                self.master.after(0, lambda: self.print_to_console(
                    f"[ERROR] Unexpected error during decryption: {e}", "error"
                ))
                self.master.after(0, lambda: messagebox.showerror(
                    "Error",
                    f"An error occurred during decryption:\n\n{e}"
                ))
        
        # Start decryption in background thread
        threading.Thread(target=decrypt_thread, daemon=True).start()

    def update_ui_for_process(self, processing):
        """Updates UI state for start/stop of the merge process."""
        state = tk.DISABLED if processing else tk.NORMAL
        self.start_btn.config(state=tk.DISABLED if processing else tk.NORMAL)
        self.pause_btn.config(state=tk.NORMAL if processing else tk.DISABLED)
        # Update pause button text based on state
        if processing:
            self.pause_btn.config(text="⏸")
        else:
            self.pause_btn.config(text="⏸")
        self.stop_btn.config(state=tk.NORMAL if processing else tk.DISABLED)
        
        # Add PDF button is now in toolbar
        if hasattr(self, 'add_btn'):
            self.add_btn.config(state=state)
        self.pdf_listbox.config(state=state)
        self.select_folder_btn.config(state=state)
        self.remove_timestamps_checkbox.config(state=state)
        self.remove_images_checkbox.config(state=state)
        self.remove_pii_checkbox.config(state=state)
        self.custom_pii_entry.config(state=state)
        self.custom_pii_label.config(state=state)
        # New: Disable split controls during processing
        self.split_by_words_checkbox.config(state=state)
        self.split_word_count_entry.config(state=state)
        self.split_word_count_label.config(state=state)
        # New: Disable markdown controls during processing
        self.generate_markdown_checkbox.config(state=state)
        self.simple_markdown_checkbox.config(state=state)
        self.use_gpu_checkbox.config(state=state)
        # Always allow changing models folder (unless currently downloading)
        if self.preload_models_btn.cget('text') != "Downloading...":
            self.preload_models_btn.config(state=state)

        if not processing:
            self._update_pii_field_visibility()
            self._update_split_field_visibility()
            self._update_markdown_controls_visibility()

        self.on_listbox_select(None)

    def start_merge(self):
        """Starts the merge process in a new thread."""
        global merge_thread, merge_running
        if merge_running:
            self.print_to_console("Merge process is already running.", "info")
            return
        if not self.pdf_files:
            messagebox.showerror("No Files", "Please add PDF files to the list first.")
            return
        
        # New: Validate split word count if enabled
        if self.split_by_words_var.get():
            try:
                word_limit = int(self.split_word_count_var.get())
                if word_limit <= 0:
                    messagebox.showerror("Invalid Input", "Word count for splitting must be a positive number.")
                    return
            except ValueError:
                messagebox.showerror("Invalid Input", "Word count for splitting must be a valid number.")
                return

        merge_stop_event.clear()
        merge_pause_event.clear()
        merge_running = True
        
        self.update_ui_for_process(processing=True)
        self.print_to_console("Starting PDF merge process...", "info")
        
        merge_thread = threading.Thread(target=self._merge_pdfs_threaded, daemon=True)
        merge_thread.start()

    def pause_merge(self):
        """Pauses or resumes the merge process."""
        global merge_paused
        if not merge_running: return
        merge_paused = not merge_paused
        if merge_paused:
            merge_pause_event.set()
            self.pause_btn.config(text="▶")  # Resume icon when paused
            self.print_to_console("Merge process paused.", "info")
        else:
            merge_pause_event.clear()
            self.pause_btn.config(text="⏸")  # Pause icon when running
            self.print_to_console("Merge process resumed.", "info")

    def stop_merge(self):
        """Stops the merge process."""
        if merge_running:
            merge_stop_event.set()
            self.print_to_console("Stopping merge process...", "info")

    def _get_output_filepath(self, counter=None, extension=None):
        """Generates a unique output filename, with an optional counter for splitting."""
        # Use custom filename if provided, otherwise use default
        if self.output_filename_var.get().strip():
            base_filename = self.output_filename_var.get().strip()
            # Remove extension if provided
            base, _ = os.path.splitext(base_filename)
        else:
            base, _ = os.path.splitext(DEFAULT_OUTPUT_FILENAME)

        # Determine extension based on output type
        if extension is None:
            output_type = self.output_file_type_var.get().lower()
            extension_map = {
                'pdf': '.pdf',
                'odt': '.odt',
                'docx': '.docx',
                'txt': '.txt',
                'rtf': '.rtf',
                'epub': '.epub',
                'md': '.md'
            }
            extension = extension_map.get(output_type, '.pdf')

        # New: Handle numbered files for splitting
        if counter is not None and counter > 1:
            base = f"{base}{counter}"

        output_filepath = os.path.join(self.output_folder, f"{base}{extension}")

        # Ensure the first file is also unique if it exists
        if counter is None or counter == 1:
            path_template = os.path.join(self.output_folder, base)
            file_counter = 1
            while os.path.exists(output_filepath):
                file_counter += 1
                output_filepath = f"{path_template}{file_counter}{extension}"

        return output_filepath

    def _scrub_pii_from_doc(self, doc):
        """
        Finds and applies redactions for PII in a PyMuPDF document object.
        Uses both Regex patterns and a custom string list for redaction.
        """
        self.print_to_console("    - Starting PII scrubbing...", "progress")
        total_redactions = 0
        
        custom_strings_raw = self.custom_pii_var.get()
        if custom_strings_raw:
            custom_strings = [s.strip() for s in custom_strings_raw.split(',') if s.strip()]
            if custom_strings:
                self.print_to_console(f"    - Searching for {len(custom_strings)} custom string(s)...", "progress")
                for page_num, page in enumerate(doc, 1):
                    for custom_string in custom_strings:
                        matches = page.search_for(custom_string)
                        if matches:
                            self.print_to_console(f"    - Found '{custom_string}' {len(matches)} time(s) on page {page_num}. Marking for redaction.", "progress")
                            for inst in matches:
                                page.add_redact_annot(inst, text=" ", fill=(0, 0, 0))
                            total_redactions += len(matches)

        pii_patterns = {
            "FULL_NAME": r'\b[A-Z]{4,}\s[A-Z]{4,}\b',
            "STREET_ADDRESS": r'\b\d{1,5}\s(?:[A-Z0-9]+\s?)+(?:STREET|ST|AVENUE|AVE|ROAD|RD|LANE|LN|DRIVE|DR|COURT|CT|PLACE|PL|BOULEVARD|BLVD)\b',
            "CITY_STATE_ZIP": r'\b[A-Z\s]+,\s[A-Z]{2}\s\d{5}(?:-\d{4})?\b',
            "ACCOUNT_NUMBER": r'\b\d{5}-\d{5}(?:-\d)?\b',
            "ID_NUMBER": r'\b\d{8,19}\b',
            "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
        }
        
        for pii_type, pattern in pii_patterns.items():
            for page_num, page in enumerate(doc, 1):
                try:
                    matches = page.search_for(pattern)
                    if matches:
                        self.print_to_console(f"    - Found {len(matches)} potential '{pii_type}' on page {page_num}. Marking for redaction.", "progress")
                        for inst in matches:
                            page.add_redact_annot(inst, text=" ", fill=(0, 0, 0))
                        total_redactions += len(matches)
                except Exception as e:
                    self.print_to_console(f"    - Error searching for PII on page {page_num}: {e}", "error")

        if total_redactions > 0:
            self.print_to_console(f"    - Applying {total_redactions} total redactions...", "progress")
            for page in doc:
                page.apply_redactions(images=fitz.PDF_REDACT_IMAGE_PIXELS)
            self.print_to_console("    - PII scrubbing complete.", "success")
        else:
            self.print_to_console("    - No PII matching the defined patterns or custom strings was found.", "warning")
            
        return doc

    def _convert_pdf_to_markdown(self, pdf_path):
        """Converts a PDF file to markdown using marker-pdf library."""
        try:
            self.print_to_console(f"    - Converting to markdown: {os.path.basename(pdf_path)}", "progress")
            
            # Set up environment variables each time (thread-safe)
            os.makedirs(self.models_directory, exist_ok=True)
            os.environ['TORCH_HOME'] = self.models_directory
            os.environ['HF_HOME'] = self.models_directory
            os.environ['TRANSFORMERS_CACHE'] = self.models_directory
            # IMPORTANT: Set Surya model cache directory
            os.environ['MODEL_CACHE_DIR'] = self.models_directory
            
            # CRITICAL: Disable multiprocessing to prevent process pool errors in frozen executable
            os.environ['MARKER_NO_MULTIPROCESSING'] = '1'
            os.environ['OMP_NUM_THREADS'] = '1'
            os.environ['MKL_NUM_THREADS'] = '1'
            
            # Disable tqdm completely before any imports
            os.environ['TQDM_DISABLE'] = '1'
            
            # Import and patch tqdm first, before any marker imports
            import sys
            import io
            
            # Redirect stdout/stderr temporarily to catch tqdm issues
            old_stdout = sys.stdout
            old_stderr = sys.stderr
            
            try:
                # Import tqdm and completely disable it
                import tqdm
                
                # Create a dummy tqdm class that does nothing
                class DummyTqdm:
                    def __init__(self, *args, **kwargs):
                        pass
                    def __enter__(self):
                        return self
                    def __exit__(self, *args):
                        pass
                    def update(self, *args):
                        pass
                    def close(self):
                        pass
                    def set_description(self, *args):
                        pass
                
                # Replace tqdm with dummy
                tqdm.tqdm = DummyTqdm
                
                # Now import marker modules
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                from marker.output import text_from_rendered
                import torch
                
            finally:
                # Restore stdout/stderr
                sys.stdout = old_stdout
                sys.stderr = old_stderr
            
            self.print_to_console("    - Creating marker-pdf models...", "progress")
            
            # Create models (exactly like working test script)
            models = create_model_dict()
            self.print_to_console(f"    - Models created: {list(models.keys())}", "info")
            
            # Determine device and move models
            if self.use_gpu_var.get() and torch.cuda.is_available():
                device = "cuda"
                self.print_to_console(f"    - Using GPU: {torch.cuda.get_device_name(0)}", "info")
                # Move models to GPU
                for key, model in models.items():
                    if hasattr(model, 'to'):
                        models[key] = model.to(device)
                        self.print_to_console(f"    - Moved {key} to GPU", "info")
            else:
                device = "cpu"
                self.print_to_console("    - Using CPU for processing", "info")
            
            self.print_to_console("    - Models created successfully.", "success")
            
            # Create converter
            self.print_to_console("    - Creating PdfConverter...", "progress")
            converter = PdfConverter(artifact_dict=models)
            self.print_to_console("    - PdfConverter created successfully.", "success")
            
            self.print_to_console(f"    - Processing PDF with marker-pdf...", "progress")
            try:
                rendered = converter(pdf_path)
                self.print_to_console(f"    - PDF processing completed, type: {type(rendered)}", "success")
            except Exception as conv_error:
                self.print_to_console(f"    - Error during PDF conversion: {conv_error}", "error")
                raise conv_error
            
            self.print_to_console(f"    - Extracting markdown text...", "progress")
            try:
                full_text, _, images = text_from_rendered(rendered)
                self.print_to_console(f"    - Text extraction completed: {len(full_text)} characters", "success")
            except Exception as text_error:
                self.print_to_console(f"    - Error during text extraction: {text_error}", "error")
                raise text_error
            
            # Apply timestamp removal if enabled
            if self.remove_timestamps_var.get():
                timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
                full_text = re.sub(timestamp_regex, '', full_text)
            
            # Apply custom PII removal if enabled
            if self.remove_pii_var.get():
                full_text = self._scrub_pii_from_text(full_text)
            
            return full_text
            
        except Exception as e:
            self.print_to_console(f"    - Error converting {os.path.basename(pdf_path)} to markdown: {e}", "error")
            return None

    def _scrub_pii_from_text(self, text):
        """Scrubs PII from text content using regex patterns."""
        try:
            # Apply custom string removal
            custom_strings_raw = self.custom_pii_var.get()
            if custom_strings_raw:
                custom_strings = [s.strip() for s in custom_strings_raw.split(',') if s.strip()]
                for custom_string in custom_strings:
                    text = text.replace(custom_string, "[REDACTED]")
            
            # Apply PII pattern removal
            pii_patterns = {
                "FULL_NAME": r'\b[A-Z]{4,}\s[A-Z]{4,}\b',
                "STREET_ADDRESS": r'\b\d{1,5}\s(?:[A-Z0-9]+\s?)+(?:STREET|ST|AVENUE|AVE|ROAD|RD|LANE|LN|DRIVE|DR|COURT|CT|PLACE|PL|BOULEVARD|BLVD)\b',
                "CITY_STATE_ZIP": r'\b[A-Z\s]+,\s[A-Z]{2}\s\d{5}(?:-\d{4})?\b',
                "ACCOUNT_NUMBER": r'\b\d{5}-\d{5}(?:-\d)?\b',
                "ID_NUMBER": r'\b\d{8,19}\b',
                "EMAIL": r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b',
            }
            
            for pii_type, pattern in pii_patterns.items():
                text = re.sub(pattern, "[REDACTED]", text, flags=re.IGNORECASE)
            
            return text
            
        except Exception as e:
            self.print_to_console(f"    - Error scrubbing PII from text: {e}", "error")
            return text

    def _convert_pdf_to_markdown_main_thread(self, pdf_path):
        """Converts a PDF file to markdown using marker-pdf library with proper GPU support"""
        try:
            self.print_to_console(f"    - Converting to markdown: {os.path.basename(pdf_path)}", "progress")
            
            # Set up environment (exactly like working test_marker.py)
            os.makedirs(self.models_directory, exist_ok=True)
            os.environ['TORCH_HOME'] = self.models_directory
            os.environ['HF_HOME'] = self.models_directory
            os.environ['TRANSFORMERS_CACHE'] = self.models_directory
            # IMPORTANT: Set Surya model cache directory
            os.environ['MODEL_CACHE_DIR'] = self.models_directory
            
            # CRITICAL: Disable multiprocessing to prevent process pool errors in frozen executable
            os.environ['MARKER_NO_MULTIPROCESSING'] = '1'
            os.environ['OMP_NUM_THREADS'] = '1'
            os.environ['MKL_NUM_THREADS'] = '1'
            
            # Set device for Surya models (the correct way!)
            import torch
            if self.use_gpu_var.get() and torch.cuda.is_available():
                os.environ['TORCH_DEVICE'] = 'cuda'
                self.print_to_console(f"    - Using GPU: {torch.cuda.get_device_name(0)}", "info")
            else:
                os.environ['TORCH_DEVICE'] = 'cpu'
                self.print_to_console("    - Using CPU for processing", "info")
            
            # Set device environment variable before importing
            if self.use_gpu_var.get() and torch.cuda.is_available():
                os.environ['TORCH_DEVICE'] = 'cuda'
            else:
                os.environ['TORCH_DEVICE'] = 'cpu'
            
            self.print_to_console("    - Importing marker modules...", "progress")
            from marker.converters.pdf import PdfConverter
            from marker.models import create_model_dict
            from marker.output import text_from_rendered
            self.print_to_console("    - Imports successful", "success")

            self.print_to_console("    - Creating model dictionary...", "progress")
            models = create_model_dict()
            self.print_to_console(f"    - Models created: {list(models.keys())}", "success")
            
            # Check that models are properly loaded
            for key, model in models.items():
                if model is None:
                    self.print_to_console(f"    - WARNING: {key} is None!", "warning")
                else:
                    self.print_to_console(f"    - {key}: {type(model).__name__} loaded", "info")
            
            self.print_to_console("    - Creating PdfConverter...", "progress")
            converter = PdfConverter(artifact_dict=models)
            self.print_to_console(f"    - Converter created: {type(converter)}", "success")
            
            self.print_to_console(f"    - Processing: {os.path.basename(pdf_path)}", "progress")
            rendered = converter(pdf_path)
            self.print_to_console(f"    - Rendered result: {type(rendered)}", "success")
            
            text, _, images = text_from_rendered(rendered)
            self.print_to_console(f"    - Text extracted: {len(text)} characters", "success")
            
            # Apply timestamp removal if enabled
            if self.remove_timestamps_var.get():
                timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
                text = re.sub(timestamp_regex, '', text)
            
            # Apply custom PII removal if enabled
            if self.remove_pii_var.get():
                text = self._scrub_pii_from_text(text)
            
            return text
            
        except Exception as e:
            self.print_to_console(f"    - Error converting {os.path.basename(pdf_path)} to markdown: {e}", "error")
            import traceback
            traceback.print_exc()
            return None

    def _convert_pdf_to_markdown_threaded(self, pdf_path):
        """Converts a PDF file to markdown using marker-pdf library - runs in background thread with output capture"""
        try:
            self.print_to_console(f"    - Converting to markdown: {os.path.basename(pdf_path)}", "progress")
            
            # Set up environment (same as working approach)
            os.makedirs(self.models_directory, exist_ok=True)
            os.environ['TORCH_HOME'] = self.models_directory
            os.environ['HF_HOME'] = self.models_directory
            os.environ['TRANSFORMERS_CACHE'] = self.models_directory
            # IMPORTANT: Set Surya model cache directory
            os.environ['MODEL_CACHE_DIR'] = self.models_directory
            
            # CRITICAL: Disable multiprocessing to prevent process pool errors in frozen executable
            os.environ['MARKER_NO_MULTIPROCESSING'] = '1'
            os.environ['OMP_NUM_THREADS'] = '1'
            os.environ['MKL_NUM_THREADS'] = '1'
            
            # Set device for Surya models
            import torch
            if self.use_gpu_var.get() and torch.cuda.is_available():
                os.environ['TORCH_DEVICE'] = 'cuda'
                self.print_to_console(f"    - Using GPU: {torch.cuda.get_device_name(0)}", "info")
            else:
                os.environ['TORCH_DEVICE'] = 'cpu'
                self.print_to_console("    - Using CPU for processing", "info")
            
            # Capture stdout/stderr to redirect marker progress to our console
            import sys
            import io
            from contextlib import redirect_stdout, redirect_stderr
            
            # Create custom output handlers
            class ThreadSafeConsoleCapture:
                def __init__(self, console_func, master, tag="progress"):
                    self.console_func = console_func
                    self.master = master
                    self.tag = tag
                    self.buffer = ""
                
                def write(self, text):
                    if text.strip():
                        # Schedule GUI update in main thread
                        self.master.after(0, lambda: self.console_func(f"    - {text.strip()}", self.tag))
                
                def flush(self):
                    pass
            
            stdout_capture = ThreadSafeConsoleCapture(self.print_to_console, self.master, "progress")
            stderr_capture = ThreadSafeConsoleCapture(self.print_to_console, self.master, "warning")
            
            self.print_to_console("    - Importing marker modules...", "progress")
            
            # Set environment variables before importing
            os.environ['TORCH_DEVICE'] = 'cuda' if (self.use_gpu_var.get() and torch.cuda.is_available()) else 'cpu'
            
            # Import and run with output capture
            with redirect_stdout(stdout_capture), redirect_stderr(stderr_capture):
                from marker.converters.pdf import PdfConverter
                from marker.models import create_model_dict
                from marker.output import text_from_rendered
                
                self.print_to_console("    - Creating model dictionary...", "progress")
                models = create_model_dict()
                self.print_to_console(f"    - Models created: {list(models.keys())}", "success")
                
                # Check that models are properly loaded
                for key, model in models.items():
                    if model is None:
                        self.print_to_console(f"    - WARNING: {key} is None!", "warning")
                
                self.print_to_console("    - Creating PdfConverter...", "progress")
                converter = PdfConverter(artifact_dict=models)
                self.print_to_console(f"    - Converter created successfully", "success")
                
                self.print_to_console(f"    - Processing PDF with marker-pdf...", "progress")
                rendered = converter(pdf_path)
                self.print_to_console(f"    - PDF processing completed", "success")
                
                self.print_to_console(f"    - Extracting markdown text...", "progress")
                text, _, images = text_from_rendered(rendered)
                self.print_to_console(f"    - Text extracted: {len(text)} characters", "success")
            
            # Apply timestamp removal if enabled
            if self.remove_timestamps_var.get():
                timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
                text = re.sub(timestamp_regex, '', text)
            
            # Apply custom PII removal if enabled
            if self.remove_pii_var.get():
                text = self._scrub_pii_from_text(text)
            
            return text
            
        except Exception as e:
            self.print_to_console(f"    - Error converting {os.path.basename(pdf_path)} to markdown: {e}", "error")
            return None

    def _merge_pdfs_threaded(self):
        """The core multi-format file processing and merging logic that runs in a thread."""
        global merge_running, merge_paused

        try:
            # --- Stage 1: Extract text from all files ---
            total_files = len(self.pdf_files)
            merged_text = ""

            for i, file_path in enumerate(self.pdf_files):
                if merge_stop_event.is_set(): break
                while merge_pause_event.is_set(): time.sleep(0.1)

                self.print_to_console(f"Processing '{os.path.basename(file_path)}' ({i+1}/{total_files})...", "progress")

                try:
                    # Extract text from file
                    text = self._extract_text_from_file(file_path)

                    # Apply text processing options
                    if self.remove_timestamps_var.get():
                        timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
                        text = re.sub(timestamp_regex, '', text)

                    # Apply PII scrubbing
                    if self.remove_pii_var.get():
                        text = self._scrub_pii_from_text(text)

                    # Add to merged text
                    merged_text += text + "\n\n"

                except Exception as e:
                    self.print_to_console(f"  Error processing '{os.path.basename(file_path)}': {e}. Skipping.", "error")
                    continue

                progress_percent = int(((i + 1) / total_files) * 100)
                self.print_to_console(f"  Processing progress: {progress_percent}%", "progress")

            if merge_stop_event.is_set():
                self.print_to_console("Process stopped during file processing.", "warning")
                raise SystemExit()

            if not merged_text.strip():
                self.print_to_console("No content was successfully processed to merge.", "warning")
                raise SystemExit()

            self.print_to_console("All files processed. Starting final merge...", "progress")

            # --- Stage 2: Generate output file(s) ---
            saved_files = []

            if self.split_by_words_var.get():
                # Split by words
                try:
                    words_per_file = int(self.split_word_count_var.get())
                except ValueError:
                    words_per_file = 10000
                    self.print_to_console(f"Invalid word count, using default: {words_per_file}", "warning")

                words = merged_text.split()
                file_counter = 1
                while words:
                    if merge_stop_event.is_set(): break

                    chunk_words = words[:words_per_file]
                    words = words[words_per_file:]
                    chunk_text = ' '.join(chunk_words)

                    output_filepath = self._get_output_filepath(counter=file_counter)
                    self._generate_output_file(chunk_text, output_filepath)
                    saved_files.append(output_filepath)
                    self.print_to_console(f"Saved part {file_counter}: {os.path.basename(output_filepath)}", "success")
                    file_counter += 1
            else:
                # Standard merging (single file)
                output_filepath = self._get_output_filepath()
                self._generate_output_file(merged_text, output_filepath)
                saved_files.append(output_filepath)
                self.print_to_console(f"Merge completed successfully: {os.path.basename(output_filepath)}", "success")

            # --- Generate Markdown if output type is MD and advanced mode selected ---
            if self.output_file_type_var.get() == "MD" and self.markdown_type_var.get() == "advanced" and saved_files:
                self.print_to_console("Advanced markdown conversion not yet implemented for multi-format merging.", "warning")

        except SystemExit: # Graceful exit on stop
             self.print_to_console("Merge process was stopped by user. No file saved.", "info")
        except Exception as e:
            self.print_to_console(f"An unexpected error occurred during the merge process: {e}", "error")
            import traceback
            traceback.print_exc()
        finally:
            merge_running = False
            merge_paused = False
            self.master.after(0, lambda: self.update_ui_for_process(processing=False))

    def _merge_standard(self, temp_files):
        """Merges all temp files into a single output PDF. Returns list of saved PDF paths."""
        final_merged_doc = fitz.open()
        for temp_path in temp_files:
            try:
                with fitz.open(temp_path) as temp_doc:
                    final_merged_doc.insert_pdf(temp_doc)
            except Exception as e:
                self.print_to_console(f"Could not merge temp file {os.path.basename(temp_path)}: {e}", "error")

        saved_pdfs = []
        if len(final_merged_doc) > 0:
            output_filepath = self._get_output_filepath()
            final_merged_doc.save(output_filepath)
            self.print_to_console(f"Successfully merged PDFs to: {output_filepath}", "success")
            saved_pdfs.append(output_filepath)
        else:
            self.print_to_console("Final document is empty after merge attempts.", "warning")

        final_merged_doc.close()
        return saved_pdfs

    def _merge_with_splitting(self, temp_files):
        """Merges temp files into multiple PDFs, split by word count. Returns list of saved PDF paths."""
        try:
            word_limit = int(self.split_word_count_var.get())
            self.print_to_console(f"Splitting output into files of approximately {word_limit} words.", "info")
        except ValueError:
            self.print_to_console("Invalid word count for splitting. Aborting.", "error")
            return []
            
        output_doc = fitz.open()
        current_word_count = 0
        file_counter = 1
        saved_pdfs = []

        for temp_path in temp_files:
            if merge_stop_event.is_set(): break
            try:
                with fitz.open(temp_path) as temp_doc:
                    for page in temp_doc:
                        if merge_stop_event.is_set(): break
                        page_text = page.get_text("text")
                        page_word_count = self._count_words(page_text)
                        
                        # If adding this page exceeds the limit, save the current doc first
                        if current_word_count + page_word_count > word_limit and current_word_count > 0:
                            output_filepath = self._get_output_filepath(file_counter)
                            output_doc.save(output_filepath)
                            self.print_to_console(f"Saved split file: {os.path.basename(output_filepath)} ({current_word_count} words)", "success")
                            saved_pdfs.append(output_filepath)
                            
                            output_doc.close()
                            output_doc = fitz.open()
                            current_word_count = 0
                            file_counter += 1
                        
                        # Add the page to the current output doc
                        output_doc.insert_pdf(temp_doc, from_page=page.number, to_page=page.number)
                        current_word_count += page_word_count
            except Exception as e:
                self.print_to_console(f"Error during splitting of {os.path.basename(temp_path)}: {e}", "error")

        # Save the last remaining document if it has content
        if not merge_stop_event.is_set() and len(output_doc) > 0:
            output_filepath = self._get_output_filepath(file_counter)
            output_doc.save(output_filepath)
            self.print_to_console(f"Saved final split file: {os.path.basename(output_filepath)} ({current_word_count} words)", "success")
            saved_pdfs.append(output_filepath)

        if saved_pdfs:
            self.print_to_console(f"Successfully created {len(saved_pdfs)} split PDF files.", "success")
        else:
            self.print_to_console("No split files were generated.", "warning")
            
        output_doc.close()
        return saved_pdfs

    def _generate_markdown_output_main_thread(self):
        """Generates a combined markdown file from all PDF files - runs in main thread."""
        try:
            self.print_to_console("Starting markdown generation in main thread...", "progress")
            
            combined_markdown = ""
            total_files = len(self.pdf_files)
            
            for i, pdf_path in enumerate(self.pdf_files):
                self.print_to_console(f"Converting to markdown: {os.path.basename(pdf_path)} ({i+1}/{total_files})", "progress")
                
                # Add file header
                filename = os.path.basename(pdf_path)
                combined_markdown += f"\n\n# {filename}\n\n"
                combined_markdown += f"*Source: {filename}*\n\n"
                combined_markdown += "---\n\n"
                
                # Convert PDF to markdown (now in main thread)
                markdown_content = self._convert_pdf_to_markdown_main_thread(pdf_path)
                if markdown_content:
                    combined_markdown += markdown_content
                    combined_markdown += "\n\n"
                else:
                    combined_markdown += "*[Error: Could not convert this PDF to markdown]*\n\n"
            
            # Save combined markdown file
            if combined_markdown.strip():
                markdown_filepath = self._get_output_filepath(extension=".md")
                
                try:
                    with open(markdown_filepath, 'w', encoding='utf-8') as f:
                        # Add document header
                        f.write("# Combined PDF Content\n\n")
                        f.write(f"*Generated from {len(self.pdf_files)} PDF files*\n\n")
                        f.write("---\n")
                        f.write(combined_markdown)
                    
                    self.print_to_console(f"[OK] Successfully generated markdown: {os.path.basename(markdown_filepath)}", "success")
                    
                except Exception as e:
                    self.print_to_console(f"[ERROR] Failed to save markdown file: {e}", "error")
            else:
                self.print_to_console("[WARNING] No markdown content generated.", "warning")
                
        except Exception as e:
            self.print_to_console(f"[ERROR] Unexpected error during markdown generation: {e}", "error")

    def _generate_markdown_output(self):
        """Generates a combined markdown file from all PDF files - runs in background thread."""
        try:
            self.print_to_console("Starting markdown generation...", "progress")
            
            combined_markdown = ""
            total_files = len(self.pdf_files)
            
            for i, pdf_path in enumerate(self.pdf_files):
                if merge_stop_event.is_set():
                    break
                    
                self.print_to_console(f"Converting to markdown: {os.path.basename(pdf_path)} ({i+1}/{total_files})", "progress")
                
                # Add file header
                filename = os.path.basename(pdf_path)
                combined_markdown += f"\n\n# {filename}\n\n"
                combined_markdown += f"*Source: {filename}*\n\n"
                combined_markdown += "---\n\n"
                
                # Convert PDF to markdown (now with proper thread-safe output capture)
                markdown_content = self._convert_pdf_to_markdown_threaded(pdf_path)
                if markdown_content:
                    combined_markdown += markdown_content
                    combined_markdown += "\n\n"
                else:
                    combined_markdown += "*[Error: Could not convert this PDF to markdown]*\n\n"
            
            if merge_stop_event.is_set():
                self.print_to_console("Markdown generation stopped by user.", "warning")
                return
            
            # Save combined markdown file
            if combined_markdown.strip():
                markdown_filepath = self._get_output_filepath(extension=".md")
                
                try:
                    with open(markdown_filepath, 'w', encoding='utf-8') as f:
                        # Add document header
                        f.write("# Combined PDF Content\n\n")
                        f.write(f"*Generated from {len(self.pdf_files)} PDF files*\n\n")
                        f.write("---\n")
                        f.write(combined_markdown)
                    
                    self.print_to_console(f"[OK] Successfully generated markdown: {os.path.basename(markdown_filepath)}", "success")
                    
                except Exception as e:
                    self.print_to_console(f"[ERROR] Failed to save markdown file: {e}", "error")
            else:
                self.print_to_console("[WARNING] No markdown content generated.", "warning")
                
        except Exception as e:
            self.print_to_console(f"[ERROR] Unexpected error during markdown generation: {e}", "error")

    def _convert_merged_pdf_to_markdown(self, pdf_path):
        """Converts a single merged PDF file to markdown. Much more efficient than converting each original file."""
        try:
            if merge_stop_event.is_set():
                return
            
            pdf_basename = os.path.basename(pdf_path)
            self.print_to_console(f"Converting merged PDF to markdown: {pdf_basename}", "progress")
            
            # Choose conversion method based on user setting
            if self.simple_markdown_var.get():
                # Fast text extraction without OCR
                markdown_content = self._convert_pdf_to_markdown_simple(pdf_path)
            else:
                # Advanced OCR-based conversion with GPU
                markdown_content = self._convert_pdf_to_markdown_threaded(pdf_path)
            
            if not markdown_content:
                self.print_to_console(f"[WARNING] Failed to convert {pdf_basename} to markdown", "warning")
                return
            
            # Generate markdown output filename based on PDF filename
            base_name = os.path.splitext(pdf_path)[0]
            markdown_filepath = base_name + ".md"
            
            # Save the markdown file
            try:
                with open(markdown_filepath, 'w', encoding='utf-8') as f:
                    # Add document header
                    f.write(f"# {pdf_basename} - Markdown Export\n\n")
                    if self.simple_markdown_var.get():
                        f.write(f"*Converted from merged PDF using PyMuPDF4LLM (text extraction)*\n\n")
                    else:
                        f.write(f"*Converted from merged PDF using marker-pdf (OCR)*\n\n")
                    f.write("---\n\n")
                    f.write(markdown_content)
                
                self.print_to_console(f"[OK] Successfully generated markdown: {os.path.basename(markdown_filepath)}", "success")
                
            except Exception as e:
                self.print_to_console(f"[ERROR] Failed to save markdown file: {e}", "error")
                
        except Exception as e:
            self.print_to_console(f"[ERROR] Unexpected error during markdown conversion: {e}", "error")
    
    def _convert_pdf_to_markdown_simple(self, pdf_path):
        """Converts a PDF to markdown using PyMuPDF4LLM (fast, no OCR, extracts existing text)."""
        try:
            self.print_to_console(f"    - Using PyMuPDF4LLM for fast text extraction (no OCR)...", "progress")
            
            # Import pymupdf4llm
            try:
                import pymupdf4llm
            except ImportError:
                self.print_to_console(f"[ERROR] pymupdf4llm not installed. Install with: pip install pymupdf4llm", "error")
                return None
            
            # Convert PDF to markdown using PyMuPDF4LLM
            start_time = time.time()
            md_text = pymupdf4llm.to_markdown(pdf_path)
            elapsed = time.time() - start_time
            
            self.print_to_console(f"    - Conversion completed in {elapsed:.2f} seconds", "success")
            self.print_to_console(f"    - Extracted {len(md_text)} characters", "info")
            
            # Apply timestamp removal if enabled
            if self.remove_timestamps_var.get():
                timestamp_regex = r'\[(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\s*-->\s*(?:(?:\d{2}:)?\d{2}:\d{2}\.\d{3})\]\s*'
                md_text = re.sub(timestamp_regex, '', md_text)
            
            # Apply custom PII removal if enabled
            if self.remove_pii_var.get():
                md_text = self._scrub_pii_from_text(md_text)
            
            return md_text
            
        except Exception as e:
            self.print_to_console(f"    - Error converting {os.path.basename(pdf_path)} with PyMuPDF4LLM: {e}", "error")
            import traceback
            traceback.print_exc()
            return None

def main():
    """Main function to create and run the Tkinter application."""
    # CRITICAL: Freeze support for multiprocessing in frozen executables
    # This prevents the "process pool terminated" error and multiple instances spawning
    multiprocessing.freeze_support()
    
    root = tk.Tk()
    app = PDFMergerApp(root)
    
    def on_closing():
        if merge_running:
            merge_stop_event.set()
            if merge_thread:
                merge_thread.join(timeout=2) # Give thread time to stop
        app.save_settings() # Ensure settings are saved on close
        root.destroy()

    root.protocol("WM_DELETE_WINDOW", on_closing)
    root.mainloop()

if __name__ == "__main__":
    main()